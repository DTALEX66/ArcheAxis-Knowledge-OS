"""AXW-CAP-503 step 2: builtin converter activator wiring tests.

Proves that every builtin converter plugin's ``get_activator()`` wraps
the REAL ingestion adapter function (no invented interfaces) into a
store-callable conversion service:

- healthcheck probes pass for every real adapter module;
- ``ConversionDispatcher.get_converter()`` returns the convert callable
  only for installed AND activated plugins, and None otherwise
  (fail-closed, no silent fallback);
- ``list_active_converters()`` contains only activated items;
- re-activation is idempotent; ``activate_all_builtins()`` is
  restart-safe (re-running it on a store that already has the packs
  installed still registers the in-process converters);
- the wrapped convert runs real conversions on small real inputs
  (html/docx/pptx/xlsx) and fails closed with a clear ConverterError
  for heavy adapters (media/ocr) and for adapter-reported failures.

Heavy adapters (media transcription, OCR) only verify registration,
dispatch and fail-closed behaviour — no model downloads or tesseract
binaries are required, keeping the suite deterministic.
"""

from __future__ import annotations

import importlib.util
from pathlib import Path
from typing import Any

import pytest

from app.capability import builtin as builtin_pkg
from app.capability.builtin import (
    activate_all_builtins,
    converters_docx,
    converters_html,
    converters_media,
    converters_ocr,
    converters_pptx,
    converters_xlsx,
)
from app.capability.conversion import (
    ConversionDispatcher,
    ConverterError,
    get_active_converter,
    get_converter,
    list_active,
    list_active_converters,
    make_file_converter,
    register_active_converter,
    reset_active_converters,
)
from app.capability.store import CapabilityStore
from shared.adapter_contract import AdapterResult
from shared.plugin_manifest import load_manifest_from_mapping

REPO_ROOT = Path(__file__).resolve().parents[1]
FIXTURES = REPO_ROOT / "tests" / "fixtures"

ALL_MODULES = (
    converters_docx,
    converters_html,
    converters_media,
    converters_ocr,
    converters_pptx,
    converters_xlsx,
)
EXPECTED_PLUGIN_IDS = tuple(module.MANIFEST["plugin_id"] for module in ALL_MODULES)

_HTML_SAMPLE = (
    "<html><head><title>AXW-CAP-503 Sample</title></head>"
    "<body><article><h1>Capability Plugins</h1>"
    "<p>Pluginized conversion main chain marker phrase.</p></article></body></html>"
)


@pytest.fixture(autouse=True)
def _clean_converter_registry() -> Any:
    """Hermetic activation state: no leakage between tests."""
    reset_active_converters()
    yield
    reset_active_converters()


@pytest.fixture
def store(tmp_path: Path) -> CapabilityStore:
    return CapabilityStore(tmp_path / "capstore")


def _install(store: CapabilityStore, module: Any, *, with_activator: bool = True) -> None:
    activator = module.get_activator() if with_activator else None
    store.install_builtin(load_manifest_from_mapping(module.MANIFEST), activator)


def _make_docx(tmp_path: Path) -> Path | None:
    """Generate a minimal real .docx; fall back to the repo delivery doc."""
    try:
        from docx import Document  # python-docx (markitdown[docx])
    except ImportError:
        cand = (
            REPO_ROOT
            / "docs/architecture/imported-designs/reference-deliveries/archeaxis-2026"
            / "ArcheAxis OS Overview.docx"
        )
        return cand if cand.is_file() else None
    doc = Document()
    doc.add_heading("AXW-CAP-503 DOCX Sample", level=1)
    doc.add_paragraph("Pluginized conversion main chain works for docx.")
    out = tmp_path / "sample.docx"
    doc.save(str(out))
    return out


def _assert_conversion_or_clear_fail_closed(service: Any, source: Path, engine_probe: str) -> None:
    """Convert a real file; require success when the engine is present,
    otherwise require a clear fail-closed ConverterError."""
    if importlib.util.find_spec(engine_probe) is None:
        with pytest.raises(ConverterError):
            service.convert(source)
        return
    result = service.convert(source)
    assert isinstance(result, AdapterResult), "converter must return the real AdapterResult"
    assert result.success, f"expected successful conversion, got: {result.error}"
    assert result.content.strip(), "successful conversion must carry content"


# ── activator correctness for every builtin module ────────────────────────


def test_every_module_has_healthcheck_and_working_activator() -> None:
    for module in ALL_MODULES:
        plugin_id = module.MANIFEST["plugin_id"]

        probe = module.healthcheck()
        assert probe["ok"] is True, f"{plugin_id} healthcheck failed: {probe}"
        assert module.ADAPTER_MODULE in probe["detail"]

        activator = module.get_activator()
        assert callable(activator), f"{plugin_id} get_activator must return a callable"

        service = activator()
        assert service.plugin_id == plugin_id
        assert service.name == module.MANIFEST["name"]
        assert callable(service.convert)
        assert callable(service), "converter service itself must be callable"
        # invocation registered the service for dispatch
        assert get_active_converter(plugin_id) is service

        # direct re-invocation is idempotent (same plugin_id re-registered)
        again = activator()
        assert get_active_converter(plugin_id) is again
        assert plugin_id in list_active()


def test_every_module_registers_its_manifest_plugin_id() -> None:
    assert set(EXPECTED_PLUGIN_IDS) == {manifest.plugin_id for manifest in builtin_pkg.discover()}


# ── dispatch: installed AND activated, otherwise None (fail-closed) ───────


def test_get_converter_is_none_for_unknown_or_not_activated(store: CapabilityStore) -> None:
    dispatcher = ConversionDispatcher(store)
    assert dispatcher.get_converter("ax.builtin.converter.docx") is None  # not installed
    assert dispatcher.list_active_converters() == []
    assert list_active_converters(store) == []

    # installed WITHOUT activator → still None (installed ≠ activated)
    _install(store, converters_docx, with_activator=False)
    assert dispatcher.get_converter("ax.builtin.converter.docx") is None
    assert dispatcher.list_active_converters() == []
    assert get_converter(store, "ax.builtin.converter.docx") is None


def test_activation_makes_converter_dispatchable(store: CapabilityStore) -> None:
    _install(store, converters_docx)
    _install(store, converters_html)
    dispatcher = ConversionDispatcher(store)

    docx_service = dispatcher.get_converter("ax.builtin.converter.docx")
    assert docx_service is not None
    assert callable(docx_service.convert)
    assert get_converter(store, "ax.builtin.converter.docx") is docx_service

    html_service = dispatcher.get_converter("ax.builtin.converter.html")
    assert html_service is not None

    # installed-but-not-activated stays None (fail-closed, no fallback)
    _install(store, converters_ocr, with_activator=False)
    assert dispatcher.get_converter("ax.builtin.converter.ocr") is None

    # unknown plugin → None
    assert dispatcher.get_converter("ax.builtin.converter.nope") is None

    # list contains ONLY activated items, sorted
    assert dispatcher.list_active_converters() == [
        "ax.builtin.converter.docx",
        "ax.builtin.converter.html",
    ]
    assert list_active_converters(store) == dispatcher.list_active_converters()


def test_reactivation_is_idempotent(store: CapabilityStore) -> None:
    manifest = load_manifest_from_mapping(converters_docx.MANIFEST)
    first = store.install_builtin(manifest, converters_docx.get_activator())
    service_before = get_active_converter("ax.builtin.converter.docx")
    assert service_before is not None

    second = store.install_builtin(manifest, converters_docx.get_activator())
    assert second == first  # store record unchanged
    assert get_active_converter("ax.builtin.converter.docx") is service_before
    assert list_active() == ["ax.builtin.converter.docx"]
    assert ConversionDispatcher(store).list_active_converters() == ["ax.builtin.converter.docx"]


def test_activate_all_builtins_restart_safe(tmp_path: Path) -> None:
    root = tmp_path / "capstore"
    store1 = CapabilityStore(root)
    assert sorted(activate_all_builtins(store1)) == sorted(EXPECTED_PLUGIN_IDS)
    assert len(ConversionDispatcher(store1).list_active_converters()) == 6

    # simulated restart: same root, fresh store instance; the in-process
    # activation registry is gone → fail-closed: nothing active
    reset_active_converters()
    store2 = CapabilityStore(root)
    assert ConversionDispatcher(store2).list_active_converters() == []

    # re-activation after restart registers converters again
    assert sorted(activate_all_builtins(store2)) == sorted(EXPECTED_PLUGIN_IDS)
    dispatcher2 = ConversionDispatcher(store2)
    assert len(dispatcher2.list_active_converters()) == 6
    assert dispatcher2.get_converter("ax.builtin.converter.docx") is not None
    assert dispatcher2.get_converter("ax.builtin.converter.xlsx") is not None


# ── real conversions through the wrapped services ─────────────────────────


def test_html_converter_converts_real_file(tmp_path: Path, store: CapabilityStore) -> None:
    source = tmp_path / "article.html"
    source.write_text(_HTML_SAMPLE, encoding="utf-8")

    _install(store, converters_html)
    service = ConversionDispatcher(store).get_converter("ax.builtin.converter.html")
    assert service is not None

    if importlib.util.find_spec("trafilatura") is None:
        with pytest.raises(ConverterError):
            service.convert(source)
        return

    result = service.convert(source)
    assert isinstance(result, AdapterResult)
    assert result.success
    assert "marker phrase" in result.content

    # callable shorthand service(path) == service.convert(path)
    via_call = service(source)
    assert via_call.success
    assert via_call.content == result.content


def test_docx_converter_converts_real_file(tmp_path: Path, store: CapabilityStore) -> None:
    source = _make_docx(tmp_path)
    if source is None:
        pytest.skip("no .docx input available in this checkout")

    _install(store, converters_docx)
    service = ConversionDispatcher(store).get_converter("ax.builtin.converter.docx")
    assert service is not None
    _assert_conversion_or_clear_fail_closed(service, source, engine_probe="markitdown")


def test_pptx_converter_converts_real_file(tmp_path: Path, store: CapabilityStore) -> None:
    source = FIXTURES / "sample.pptx"
    if not source.is_file():
        pytest.skip("tests/fixtures/sample.pptx missing")

    _install(store, converters_pptx)
    service = ConversionDispatcher(store).get_converter("ax.builtin.converter.pptx")
    assert service is not None
    _assert_conversion_or_clear_fail_closed(service, source, engine_probe="pptx")


def test_xlsx_converter_converts_real_file(tmp_path: Path, store: CapabilityStore) -> None:
    source = FIXTURES / "sample.xlsx"
    if not source.is_file():
        pytest.skip("tests/fixtures/sample.xlsx missing")

    _install(store, converters_xlsx)
    service = ConversionDispatcher(store).get_converter("ax.builtin.converter.xlsx")
    assert service is not None
    _assert_conversion_or_clear_fail_closed(service, source, engine_probe="openpyxl")


def test_media_converter_registered_and_fails_closed(
    tmp_path: Path, store: CapabilityStore
) -> None:
    _install(store, converters_media)
    service = ConversionDispatcher(store).get_converter("ax.builtin.converter.media")
    assert service is not None
    assert callable(service.convert)

    # missing file → adapter reports failure → ConverterError (fail-closed)
    with pytest.raises(ConverterError, match="ax.builtin.converter.media.*file not found"):
        service.convert(tmp_path / "missing.mp3")

    # unknown option key → refused before any adapter work
    with pytest.raises(ConverterError, match="unsupported conversion option"):
        service.convert(tmp_path / "missing.mp3", {"bogus": 1})


def test_ocr_converter_registered_and_fails_closed(tmp_path: Path, store: CapabilityStore) -> None:
    _install(store, converters_ocr)
    service = ConversionDispatcher(store).get_converter("ax.builtin.converter.ocr")
    assert service is not None
    assert callable(service.convert)

    with pytest.raises(ConverterError, match="ax.builtin.converter.ocr.*file not found"):
        service.convert(tmp_path / "missing.png")

    with pytest.raises(ConverterError, match="unsupported conversion option"):
        service.convert(tmp_path / "missing.png", {"bogus": 1})


# ── option mapping onto real adapter keyword arguments (no invented API) ──


def test_ocr_lang_option_maps_to_real_adapter_kwarg(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    import app.ingestion.ocr_adapter as ocr_adapter

    seen: dict[str, Any] = {}

    def fake_convert_ocr(file_path: str, lang: str = "eng+chi_sim") -> AdapterResult:
        seen["lang"] = lang
        return AdapterResult(success=True, content="fake ocr text", engine="ocr-adapter")

    monkeypatch.setattr(ocr_adapter, "convert_ocr", fake_convert_ocr)
    service = make_file_converter(
        plugin_id="ax.builtin.converter.ocr",
        name="OCR Converter (builtin)",
        adapter_module="app.ingestion.ocr_adapter",
        adapter_function="convert_ocr",
        allowed_options=("lang",),
    )
    result = service.convert(tmp_path / "page.png", {"lang": "eng"})
    assert result.success
    assert seen["lang"] == "eng"


def test_media_work_dir_option_maps_to_real_adapter_kwarg(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    import app.ingestion.media_adapter as media_adapter

    seen: dict[str, Any] = {}

    def fake_convert_media(file_path: str, work_dir: str | None = None) -> AdapterResult:
        seen["work_dir"] = work_dir
        return AdapterResult(success=True, content="fake transcript", engine="media-adapter")

    monkeypatch.setattr(media_adapter, "convert_media", fake_convert_media)
    work = tmp_path / "work"
    work.mkdir()
    service = make_file_converter(
        plugin_id="ax.builtin.converter.media",
        name="Media Transcriber (builtin)",
        adapter_module="app.ingestion.media_adapter",
        adapter_function="convert_media",
        allowed_options=("work_dir",),
    )
    result = service.convert(tmp_path / "clip.mp3", {"work_dir": str(work)})
    assert result.success
    assert seen["work_dir"] == str(work)


# ── fail-closed guarantees of the wrapper ─────────────────────────────────


def test_wrapper_raises_on_adapter_reported_failure(monkeypatch: pytest.MonkeyPatch) -> None:
    import app.ingestion.docx_adapter as docx_adapter

    def fake_convert_docx(file_path: str) -> AdapterResult:
        return AdapterResult(
            success=False,
            content="",
            engine="docx-adapter",
            error="engine unavailable: markitdown missing",
        )

    monkeypatch.setattr(docx_adapter, "convert_docx", fake_convert_docx)
    service = make_file_converter(
        plugin_id="ax.builtin.converter.docx",
        name="DOCX Converter (builtin)",
        adapter_module="app.ingestion.docx_adapter",
        adapter_function="convert_docx",
    )
    with pytest.raises(ConverterError, match="markitdown missing"):
        service.convert("whatever.docx")


def test_wrapper_wraps_adapter_exceptions_fail_closed(monkeypatch: pytest.MonkeyPatch) -> None:
    import app.ingestion.docx_adapter as docx_adapter

    def exploding_convert_docx(file_path: str) -> AdapterResult:
        raise RuntimeError("adapter exploded")

    monkeypatch.setattr(docx_adapter, "convert_docx", exploding_convert_docx)
    service = make_file_converter(
        plugin_id="ax.builtin.converter.docx",
        name="DOCX Converter (builtin)",
        adapter_module="app.ingestion.docx_adapter",
        adapter_function="convert_docx",
    )
    with pytest.raises(ConverterError, match="adapter exploded"):
        service.convert("whatever.docx")


def test_register_active_converter_validates_input() -> None:
    with pytest.raises(ConverterError, match="plugin_id"):
        register_active_converter("", None)  # type: ignore[arg-type]
